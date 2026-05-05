from __future__ import annotations

import logging
import os
import re
from pathlib import Path
from typing import Optional

from app.config import get_settings
from app.models.models import UploadedFile
from sqlalchemy.orm import Session

logger = logging.getLogger(__name__)

_CACHE: dict[str, tuple[float, str]] = {}


def _read_docx_paragraphs(path: str) -> str:
    try:
        from docx import Document
    except ImportError as e:
        raise ImportError("python-docx is required for .docx knowledge files") from e

    doc = Document(path)
    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def load_external_knowledge_raw(path: str) -> str:
    """Load full text from a Word reference document. Cached by path + mtime."""
    if not path or not os.path.isfile(path):
        return ""
    try:
        mtime = os.path.getmtime(path)
    except OSError:
        return ""
    cached = _CACHE.get(path)
    if cached and cached[0] == mtime:
        return cached[1]
    try:
        text = _read_docx_paragraphs(path)
        text = re.sub(r"\n{3,}", "\n\n", text).strip()
        _CACHE[path] = (mtime, text)
        return text
    except Exception as e:
        logger.warning("Could not read SMC knowledge docx %s: %s", path, e)
        return ""


def get_external_knowledge_excerpt(max_chars: int = 8000) -> str:
    settings = get_settings()
    path = (settings.smc_knowledge_docx_path or "").strip()
    if not path:
        return ""
    raw = load_external_knowledge_raw(path)
    if not raw:
        return ""
    if len(raw) <= max_chars:
        return raw
    return raw[: max_chars - 20] + "\n… [truncated]"


def get_uploaded_knowledge_excerpt(
    db: Optional[Session],
    project_id: Optional[str],
    max_chars: int = 8000,
) -> str:
    """
    Get knowledge excerpt from latest uploaded .docx notes file.
    Priority: current project -> global notes.
    """
    if db is None:
        return ""

    base_q = db.query(UploadedFile).filter(UploadedFile.file_type == "notes")
    project_row = None
    if project_id:
        project_row = (
            base_q.filter(UploadedFile.project_id == project_id)
            .order_by(UploadedFile.created_at.desc())
            .all()
        )
        project_row = next(
            (r for r in project_row if Path(r.file_path).suffix.lower() == ".docx"),
            None,
        )

    global_row = None
    if project_row is None:
        global_rows = (
            base_q.filter(UploadedFile.project_id.is_(None))
            .order_by(UploadedFile.created_at.desc())
            .all()
        )
        global_row = next(
            (r for r in global_rows if Path(r.file_path).suffix.lower() == ".docx"),
            None,
        )

    row = project_row or global_row
    if not row:
        return ""
    raw = load_external_knowledge_raw(row.file_path)
    if not raw:
        return ""
    if len(raw) <= max_chars:
        return raw
    return raw[: max_chars - 20] + "\n… [truncated]"


def knowledge_doc_status() -> dict:
    settings = get_settings()
    path = (settings.smc_knowledge_docx_path or "").strip()
    if not path:
        return {"configured": False, "path": "", "loaded": False, "char_count": 0, "error": None}
    if not os.path.isfile(path):
        return {"configured": True, "path": path, "loaded": False, "char_count": 0, "error": "file_not_found"}
    raw = load_external_knowledge_raw(path)
    err = None
    if not raw:
        err = "empty_or_unreadable"
    return {
        "configured": True,
        "path": path,
        "loaded": bool(raw),
        "char_count": len(raw),
        "error": err,
    }


def get_reference_block_for_prompt(
    db: Optional[Session] = None,
    project_id: Optional[str] = None,
) -> str:
    """Single block to inject into vision / TradingView prompts (respects max_chars)."""
    settings = get_settings()
    excerpt = get_uploaded_knowledge_excerpt(db, project_id, settings.smc_knowledge_max_chars)
    if not excerpt:
        excerpt = get_external_knowledge_excerpt(settings.smc_knowledge_max_chars)
    if not excerpt:
        return ""
    return (
        "--- REFERENCE KNOWLEDGE (from Word document) ---\n"
        f"{excerpt}\n"
        "--- END REFERENCE KNOWLEDGE ---\n"
    )
